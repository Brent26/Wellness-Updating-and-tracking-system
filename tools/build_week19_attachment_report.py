from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "Week 19 - WUTS Flask Security and Local Deployment Setup.docx"

NAVY = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GREY = "F2F4F7"
TEXT = RGBColor(31, 41, 55)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=95, start=120, bottom=95, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, size=9.5, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.04
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margin(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    keep_with_next(p)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = TEXT
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = TEXT
    return p


def add_timeline_table(doc):
    timeline = [
        (
            "19 May 2026, 18:42-19:03 SAST",
            "Security and configuration hardening",
            "Implemented/verified CSRF and rate-limiting foundations, externalised configuration and branding values, removed hardcoded secret usage, and applied CSRF support to the base/login interface.",
        ),
        (
            "20 May 2026, 01:28-01:32 SAST",
            "Form CSRF repair and API test pack",
            "Added CSRF tokens to change-password and administration forms after a 400 error, and created the Postman collection for login, CSRF capture, read/write API calls, and negative input checks.",
        ),
        (
            "20 May 2026, 04:22 SAST",
            "Protected server-sent log stream",
            "Applied authentication protection to /stream/logs so unauthenticated requests are redirected to login while an authorised stream produces heartbeat pings.",
        ),
        (
            "20 May 2026, 12:54-16:14 SAST",
            "SQL Server and schema migration control",
            "Wired Flask-Migrate/Alembic into the application, recorded migration files, and baselined the existing SQL Server WUTS schema at revision ca16f031a608.",
        ),
        (
            "20 May 2026, 16:28-16:31 SAST",
            "Local test-data environment and frontend requests",
            "Created the WUTS TEST workbook/folder structure using the application's expected column headings, updated environment paths, and completed CSRF headers for frontend mutating requests.",
        ),
        (
            "21 May 2026, 09:24 SAST",
            "Technical status documentation",
            "Compiled the setup, verified API behaviour, database position, test-data sandbox, and remaining deployment/security gaps into a project status document.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.style = "Table Grid"
    widths = [Cm(3.35), Cm(4.0), Cm(10.0)]
    headers = ["Timestamp", "Major Change", "Technical Outcome"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        shade(cell, NAVY)
        set_cell_text(cell, header, bold=True, size=8.5, color="FFFFFF")
    repeat_header(table.rows[0])
    for time, change, outcome in timeline:
        cells = table.add_row().cells
        for i, value in enumerate((time, change, outcome)):
            cells[i].width = widths[i]
            set_cell_text(cells[i], value, bold=i == 1, size=8.2)
        shade(cells[0], LIGHT_BLUE)
    return table


def add_register_table(doc):
    records = [
        ("Configuration and sanitisation", ".env, .env.example, config.py, flask_config/settings.py, seed_admin.py, README.md", "Environment-driven SQL Server, folder and branding settings; placeholder configuration added; copied-company-specific details reduced."),
        ("Application startup and extensions", "app/__init__.py, app/extensions.py, requirements.txt", "CSRF, rate limiting and Flask-Migrate/Alembic dependencies and initialisation; application secret sourced from configuration."),
        ("Authentication and protected stream", "app/routes/auth.py, app/routes/stream.py", "Login controls/rate limiting and authenticated SSE log access."),
        ("User-facing CSRF handling", "app/templates/base.html, login.html, change_password.html, admin/user_form.html, admin/users.html, app/static/js/dashboard.js", "Hidden CSRF fields and X-CSRFToken headers for state-changing requests."),
        ("Migration management", "migrations/*, migrations/versions/ca16f031a608_baseline_existing_schema.py", "Alembic environment created and SQL Server schema baselined without recreating existing tables."),
        ("API testing and data sandbox", "postman_collection_wuts.json; WUTS TEST workbooks/folders; .env path values", "Repeatable local Postman testing and mapper-aligned Excel test files for periodics, biometrics, headcount and MEDIC workflows."),
        ("Documentation", "WUTS_Project_Status.docx, tools/build_wuts_status_doc.py", "Technical status summary documenting verified work and outstanding deployment gaps."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.style = "Table Grid"
    widths = [Cm(3.35), Cm(6.1), Cm(7.9)]
    for i, header in enumerate(("Work Area", "Files / Assets Touched or Created", "Change Made and Purpose")):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        shade(cell, NAVY)
        set_cell_text(cell, header, bold=True, size=8.2, color="FFFFFF")
    repeat_header(table.rows[0])
    for area, files, change in records:
        cells = table.add_row().cells
        for i, value in enumerate((area, files, change)):
            cells[i].width = widths[i]
            set_cell_text(cells[i], value, bold=i == 0, size=7.9)
        shade(cells[0], GREY)
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.text = "Weekly Project Diary | WUTS Flask application transition and local deployment preparation"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("667085")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Weekly Project Diary")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    info = doc.add_table(rows=4, cols=4)
    info.autofit = False
    info.style = "Table Grid"
    widths = [Cm(3.1), Cm(7.0), Cm(3.1), Cm(4.1)]
    metadata = [
        ("Student Name", "Brent Junior Seabelo Tshekane", "Sheet No.", "19"),
        ("Mentor's Name", "Dr Ludo Molobe", "Week Beginning", "18 May 2026"),
        ("Supervisor Name", "Dr Ludo Molobe", "", ""),
        ("Placement Company", "Debswana Diamond Company - Corporate Centre, Wellness Department", "", ""),
    ]
    for row, values in zip(info.rows, metadata):
        for i, value in enumerate(values):
            row.cells[i].width = widths[i]
            if i in (0, 2):
                shade(row.cells[i], LIGHT_BLUE)
                set_cell_text(row.cells[i], value, bold=True, size=8.8, color=NAVY)
            else:
                set_cell_text(row.cells[i], value, size=8.8)
    info.cell(3, 1).merge(info.cell(3, 3))
    set_cell_text(info.cell(3, 1), metadata[3][1], size=8.8)

    add_section_heading(doc, "Activities Planned for Week")
    add_bullet(doc, "Prepare the WUTS Flask web application for controlled local development by removing embedded environment-specific configuration and replacing it with configurable local settings.")
    add_bullet(doc, "Connect the application to a locally managed SQL Server database, establish controlled schema migration management, and create test data folders/workbooks that satisfy the application's column mappings.")
    add_bullet(doc, "Test authentication, role-controlled API access, CSRF-protected write operations and system log streaming using a repeatable Postman-based test workflow.")

    add_section_heading(doc, "Activities Actually Performed During Week")
    add_body(doc, "WUTS Flask application security and local-environment transition: Work during the week concentrated on preparing the Flask application as a controlled local development environment suitable for further development and testing. Environment-specific configuration values and paths were removed or replaced with environment-driven values, and an example environment file was introduced so future setup can be reproduced without embedding real credentials, report links or storage locations in source code. The Flask application now reads its signing secret and local resource locations from configuration instead of relying on a hardcoded application secret.")
    add_body(doc, "Authentication and request-protection testing: CSRF protection was completed across the browser workflow after a forced password-change attempt returned a 400 Bad Request stating, \"The CSRF token is missing.\" Hidden CSRF inputs were added to the login, password-change and user-management forms, and JavaScript write requests were updated to send the X-CSRFToken header. The protected log-stream route was also secured with login enforcement; testing confirmed that unauthenticated access redirects to the login screen and an authorised connection remains live through server-sent event ping messages. A Postman collection was prepared to capture CSRF tokens, maintain authenticated cookies, test read and write APIs, and exercise invalid input cases such as unknown employee-status categories and file folders.")
    add_body(doc, "Database and test-data preparation: The application was configured for a local SQL Server LocalDB instance and WUTS database. Flask-Migrate/Alembic was integrated to manage future schema changes, and the existing database schema was stamped at baseline revision ca16f031a608 rather than being recreated. A local WUTS TEST folder structure was created with mapper-compatible Excel workbooks and CSV logging assets for periodic medical records, biometrics, headcount and Medical ID processes. API checks confirmed that employee status categories, biometric statistics and test-file listings could be read from the prepared local data sources. Biometric records remain Excel-backed at this stage and have not yet been migrated into a SQL table.")

    add_section_heading(doc, "Timestamped Major Change Record")
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(5)
    note_run = note.add_run("Evidence basis: timestamps below are derived from the modified files and generated artefacts in the project folder, expressed in South Africa Standard Time (UTC+02:00).")
    note_run.italic = True
    note_run.font.name = "Arial"
    note_run.font.size = Pt(8.5)
    note_run.font.color.rgb = RGBColor.from_string("667085")
    add_timeline_table(doc)

    add_section_heading(doc, "Technical Files, Modules and Assets Touched or Created")
    add_register_table(doc)

    add_section_heading(doc, "Problems Encountered")
    add_body(doc, "Dependency and copied-environment startup issue: Running the application with the system Python raised \"ModuleNotFoundError: No module named 'flask_wtf'.\" The copied virtual environment's executable launcher also failed because it still referred to the original laptop path: \"Fatal error in launcher: Unable to create process using ... C:\\Users\\TshekaneB\\Desktop\\WUTS_FINAL_copy\\venv\\Scripts\\python.exe.\" The working procedure was changed to execute the interpreter directly through `.\\venv\\Scripts\\python.exe` and to launch Flask commands with `.\\venv\\Scripts\\python.exe -m flask ...`, avoiding the stale launcher mapping.")
    add_body(doc, "CSRF and endpoint-access failures: The browser initially returned \"400 Bad Request - The CSRF token is missing\" on the password-change page. This was resolved by inserting CSRF form fields and frontend token headers. A 404 response was also seen while testing `/stream/logs`; after the endpoint was patched and the request URL corrected, the authenticated stream connected and produced ping events. Postman testing also showed that switching username variables does not switch an existing authenticated cookie, which explained an apparent auditor-role access anomaly and reinforced the need to clear cookies or log out between role tests.")
    add_body(doc, "Database query and file-source issues: A SQL query initially failed with \"Msg 208, Level 16, State 1, Line 1 - Invalid object name 'periodic_records'.\" The issue was traced to database/schema context and resolved by querying `dbo.periodic_records` within the WUTS database. Before local file mappings had been created, API responses reporting unavailable files or folders were accepted as expected behaviour. The WUTS TEST directory and matching Excel files were then created and linked through environment configuration, allowing the relevant read APIs to return valid test results.")
    add_body(doc, "Outstanding technical concerns: The system is usable for controlled local development and API verification, but it is not yet production-ready. Remaining work includes final role-permission enforcement testing, validating and constraining file paths, session invalidation for deactivated or password-reset users, replacing Excel as mutable storage for biometric data if required, production-grade logging/monitoring, backup planning and CI/CD deployment checks.")

    add_section_heading(doc, "Mentor's Signature")
    signature = doc.add_table(rows=1, cols=4)
    signature.style = "Table Grid"
    signature.autofit = False
    sig_widths = [Cm(2.2), Cm(7.5), Cm(1.7), Cm(4.0)]
    for i, value in enumerate(("Signature", "", "Date", "")):
        signature.rows[0].cells[i].width = sig_widths[i]
        if i in (0, 2):
            shade(signature.rows[0].cells[i], LIGHT_BLUE)
            set_cell_text(signature.rows[0].cells[i], value, bold=True, size=9, color=NAVY)
        else:
            set_cell_text(signature.rows[0].cells[i], value, size=9)

    doc.core_properties.title = "Week 19 - WUTS Flask Security and Local Deployment Setup"
    doc.core_properties.subject = "Industrial attachment weekly project diary"
    doc.core_properties.author = "Brent Junior Seabelo Tshekane"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
