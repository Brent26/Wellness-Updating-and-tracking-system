from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "WUTS_Project_Status.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    p.add_run(text)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F6FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.color.rgb = RGBColor(68, 68, 68)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WUTS Project Status & Setup Summary")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Compiled on {date.today().strftime('%d %B %Y')}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph(
        "This document summarizes the WUTS setup, security cleanup, database "
        "migration work, API testing setup, and current gaps completed during "
        "the local rebuild."
    )

    add_heading(doc, "Current Run Command")
    add_code_block(
        doc,
        [
            'cd "C:\\Users\\mrtsh\\Downloads\\WUTS near prod\\WUTS_FINAL_copy"',
            ".\\venv\\Scripts\\python.exe .\\run.py",
            "http://127.0.0.1:5000/login",
        ],
    )

    add_heading(doc, "Completed Work")
    completed = [
        (
            "Secrets and redaction",
            "Company-specific names, emails, OneDrive paths, Power BI report/tenant IDs, "
            "and hardcoded secrets were replaced with environment-driven placeholders.",
        ),
        (
            "SQL Server setup",
            "The app now points to SQL Server LocalDB: (localdb)\\MSSQLLocalDB, "
            "database WUTS, using ODBC Driver 17 for SQL Server.",
        ),
        (
            "Authentication setup",
            "A wellness manager account was created/reset with a hashed password. "
            "First login forces password change.",
        ),
        (
            "CSRF fixes",
            "CSRF tokens were added to POST forms and frontend mutating requests now "
            "include X-CSRFToken headers.",
        ),
        (
            "Stream log protection",
            "/stream/logs now requires login and redirects unauthenticated requests to /login.",
        ),
        (
            "Postman testing",
            "A local Postman collection was created with login, CSRF capture, read API checks, "
            "write endpoints, and negative tests.",
        ),
        (
            "Migrations",
            "Flask-Migrate/Alembic were installed and wired. SQL Server was stamped at "
            "baseline revision ca16f031a608.",
        ),
        (
            "Test data sandbox",
            "The WUTS TEST folder, subfolders, and seed Excel/CSV files were created and "
            "wired into .env.",
        ),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    set_cell_text(header[0], "Area", True, "FFFFFF")
    set_cell_text(header[1], "Summary", True, "FFFFFF")
    shade_cell(header[0], "1F4E79")
    shade_cell(header[1], "1F4E79")
    for area, summary in completed:
        row = table.add_row().cells
        set_cell_text(row[0], area, True)
        set_cell_text(row[1], summary)

    add_heading(doc, "SQL Server State")
    add_bullets(
        doc,
        [
            "Database: WUTS on (localdb)\\MSSQLLocalDB.",
            "Current tables: users, job_runs, periodic_records, headcount_records, "
            "medic_records, conflict_logs, record_change_logs, alembic_version.",
            "Current migration baseline: ca16f031a608.",
            "Biometric data is not yet SQL-backed; it still reads/writes the biometric "
            "Excel workbook directly.",
        ],
    )

    add_heading(doc, "WUTS TEST Data Folder")
    add_code_block(
        doc,
        [
            "C:\\Users\\mrtsh\\Downloads\\OneDrive - Botswana Accountancy College\\WUTS TEST",
        ],
    )
    add_bullets(
        doc,
        [
            "Folders created: PERIODICS_PDFS, HEADCOUNT, BIOMETRICS_DATA, "
            "BIOMETRICS_DATA\\PDFS, MEDIC_IDENTIFICATION, MEDIC_IDENTIFICATION\\OUTPUT, LOGS.",
            "Master workbook: WUTS_Master_Data.xlsx with sheets DCC Medicals, New Employees, and Exits.",
            "Biometric workbook: BIOMETRICS_DATA\\Biometric_Screening_2026.xlsx with sheet 2026.",
            "Input workbooks: HEADCOUNT\\Headcount_Test_Input.xlsx and "
            "MEDIC_IDENTIFICATION\\Medic_Identification_Test_Input.xlsx.",
            "Conflict log: LOGS\\conflict_log.csv.",
            "The .env file now points the app to these test files and folders.",
        ],
    )

    add_heading(doc, "Verified API Behavior")
    add_bullets(
        doc,
        [
            "Read APIs worked as expected in Postman.",
            "/api/employee-status returns seed records for overdue, due_soon, and up_to_date categories.",
            "/api/biometric/stats reports seeded biometric coverage from the Excel file.",
            "/api/files/headcount and /api/files/medic list the test input files.",
            "Invalid inputs such as bad employee status category and unknown file folder return clean 400 JSON errors.",
        ],
    )

    add_heading(doc, "Known Gaps")
    add_bullets(
        doc,
        [
            "No biometric_records SQL table/model exists yet.",
            "Some file APIs still return full local folder paths; this should be tightened before production use.",
            "Path traversal hardening is still needed for file-processing endpoints.",
            "db.create_all() still runs at app startup; long-term schema ownership should move fully to migrations.",
            "Flask-Limiter currently uses in-memory storage, which is acceptable for local development but not production.",
            "The copied virtual environment has broken script launchers; use .\\venv\\Scripts\\python.exe -m flask instead of flask.exe.",
            "The test suite still needs stronger coverage and pytest setup verification.",
        ],
    )

    add_heading(doc, "Recommended Next Steps")
    add_bullets(
        doc,
        [
            "Add a SQL-backed biometric_records model/table and migrate it with Flask-Migrate.",
            "Refactor biometric APIs to sync/read from SQL Server, with Excel as import/export or backup only.",
            "Harden file listing and file-processing paths so only configured folders are accessible.",
            "Add is_active session validation and session invalidation on password reset/change.",
            "Replace startup db.create_all() with migration-driven schema management.",
            "Add focused pytest coverage for auth, role permissions, CSRF, and API error handling.",
        ],
    )

    add_heading(doc, "Useful Commands")
    add_code_block(
        doc,
        [
            '$env:FLASK_APP = "run.py"',
            ".\\venv\\Scripts\\python.exe -m flask db current",
            '.\\venv\\Scripts\\python.exe -m flask db migrate -m "describe change"',
            ".\\venv\\Scripts\\python.exe -m flask db upgrade",
            ".\\venv\\Scripts\\python.exe -m compileall app flask_config modules config.py run.py",
        ],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(
        "Status: Local development setup is usable with SQL Server, migrations, CSRF, "
        "secured stream logs, Postman tests, and seeded test Excel files. The next major "
        "design step is moving biometric records into SQL Server."
    )
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
